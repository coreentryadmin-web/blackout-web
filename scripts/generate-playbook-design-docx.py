#!/usr/bin/env python3
"""Generate SPX Slayer Playbook Design Word document for download."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import datetime
import os

OUT_PATH = os.environ.get(
    "PLAYBOOK_DOCX_OUT",
    "/opt/cursor/artifacts/SPX-Slayer-Playbook-Design-v1.docx",
)


def set_doc_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        h.font.bold = True


def add_title_page(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("SPX Slayer Playbook Architecture\nDesign Specification")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("BlackOut Trades — End-to-End Design Document")
    r.font.size = Pt(14)
    r.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version 1.0 | {datetime.date.today().strftime('%B %d, %Y')}\n")
    meta.add_run("Purpose: Architecture reference for implementation and AI validation\n")
    meta.add_run("Repository: blackout-web (SPX Slayer / Play Engine)")

    doc.add_paragraph()
    purpose = doc.add_paragraph(
        "This document captures the full playbook-first redesign for SPX Slayer trade alerts. "
        "It explains the current system, why weak plays feel random, the proposed architecture, "
        "twelve named playbooks, cross-tool data contracts, UI mapping, migration phases, and open "
        "decisions. Use it to validate design with Claude, ChatGPT, or engineering review before "
        "implementation."
    )
    purpose.paragraph_format.space_after = Pt(12)

    doc.add_page_break()


def add_toc_placeholder(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    sections = [
        "1. Executive Summary",
        "2. Problem Statement — Why Plays Feel Weak",
        "3. Current Architecture (As-Built)",
        "4. Target Architecture — Playbook-First",
        "5. Play State Machine",
        "6. The Twelve Playbooks (Full Catalog)",
        "7. Cross-Tool Data Plane",
        "8. Confluence vs Playbook Checklists",
        "9. Largo, Night Hawk, and Parallel Engines",
        "10. UI Mapping — Trade Alerts Panel",
        "11. Telemetry and Win-Rate Measurement",
        "12. Migration Roadmap (Shadow → Live)",
        "13. Open Design Decisions",
        "14. Validation Checklist for Other AIs",
        "15. Glossary",
    ]
    for s in sections:
        doc.add_paragraph(s, style="List Number")
    doc.add_page_break()


def add_executive_summary(doc: Document) -> None:
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "SPX Slayer today uses an additive confluence scorer (~20 weighted factors) plus generic "
        "gates and confirmations. When a BUY fires, traders see a pile of factors but not a named "
        "setup identity. That makes weak trades feel random and prevents per-pattern win-rate "
        "measurement."
    )
    doc.add_paragraph(
        "The proposed redesign routes market context through a Regime Router, matches one of "
        "twelve explicit playbooks, and drives the Trade Alerts UI from playbook state: "
        "ARMED (watch) → TRIGGERED → OPEN → MANAGING → CLOSED. Confluence becomes a checklist "
        "for the active playbook only—not a global factor soup."
    )
    doc.add_paragraph(
        "Key principle: Every visible trade must answer three questions: "
        "(1) What setup is this? (2) What must happen to enter? (3) What invalidates it?"
    )


def add_problem(doc: Document) -> None:
    doc.add_heading("2. Problem Statement — Why Plays Feel Weak", level=1)

    doc.add_heading("2.1 Symptom", level=2)
    for item in [
        "BUY alerts fire without a clear narrative traders can repeat.",
        "Watch list uses a coarse key (0dte:{direction}:{date})—not setup-specific.",
        "Failed plays look like 'confluence was high' rather than 'this pattern failed.'",
        "No playbook_id on outcomes—telemetry cannot rank patterns.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2.2 Root Cause", level=2)
    doc.add_paragraph(
        "The engine optimizes for a scalar score (confluence %) rather than discrete setup "
        "recognition. Multiple parallel philosophies exist (main play, lotto, power hour, "
        "Night Hawk nudge) without a single primary armed playbook at any moment."
    )

    doc.add_heading("2.3 Design Goal", level=2)
    doc.add_paragraph(
        "Replace 'highest score wins' with 'best eligible playbook match wins.' "
        "One primary armed playbook at a time. BUY requires a playbook-specific trigger, not "
        "merely crossing a global threshold."
    )


def add_current_arch(doc: Document) -> None:
    doc.add_heading("3. Current Architecture (As-Built)", level=1)

    doc.add_heading("3.1 Pipeline", level=2)
    pipeline = [
        "buildSpxDesk() — aggregates desk payload (spot, regime, walls, flows, technicals).",
        "computeSpxConfluence() — ~20 weighted factors → single score.",
        "evaluatePlayGates() — halt, session, tier, launch gates.",
        "evaluatePlayConfirmations() — directional confirmations.",
        "evaluateSpxPlay() — BUY / WATCH / HOLD decision.",
        "Optional Claude layer for narrative (advisory).",
    ]
    for step in pipeline:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("3.2 Poll Cadence", level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Data stream", "Interval"),
        ("Play payload (useSpxPlay)", "~3 seconds (PLAY_MS = 3_000)"),
        ("Desk pulse", "~1 second"),
        ("Flow feed", "~2 seconds"),
        ("Full desk rebuild", "~10 seconds"),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b

    doc.add_paragraph()
    doc.add_paragraph(
        "Confluence factors refresh on the play poll (~3s), not every second. UI may feel "
        "faster due to spot/desk pulse updates."
    )

    doc.add_heading("3.3 Watch Key Today", level=2)
    doc.add_paragraph(
        "Format: 0dte:{direction}:{date} — too coarse. Two different setups on the same day "
        "and direction collapse into one watch slot."
    )

    doc.add_heading("3.4 Largo Role Today", level=2)
    doc.add_paragraph(
        "Largo reads SPX play state via get_spx_play / getSpxPlayState(). It narrates and "
        "advises; it does NOT gate entries. This separation should remain unless explicitly "
        "changed."
    )


def add_target_arch(doc: Document) -> None:
    doc.add_heading("4. Target Architecture — Playbook-First", level=1)

    doc.add_heading("4.1 High-Level Flow", level=2)
    flow = [
        "Desk + cross-tool inputs arrive on schedule.",
        "Regime Router classifies session context (trend, chop, pin, power hour, etc.).",
        "Playbook Registry filters to eligible playbooks for this regime.",
        "Matcher scores each eligible playbook's preconditions.",
        "Winner becomes PRIMARY armed playbook (one at a time).",
        "State machine advances: IDLE → ARMED → TRIGGERED → OPEN → MANAGING → CLOSED.",
        "UI renders Open box, Watch box, and playbook-specific confluence checklist.",
    ]
    for step in flow:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("4.2 Core Components (Proposed)", level=2)
    components = [
        ("playbook-registry.ts", "Static definitions: id, name, regime tags, preconditions, triggers, invalidations, targets."),
        ("regime-router.ts", "Maps desk.regime + session clock + volatility to eligible playbook set."),
        ("playbook-matcher.ts", "Scores preconditions; picks primary armed playbook."),
        ("playbook-state.ts", "Persistent state machine per playbook instance."),
        ("playbook-telemetry.ts", "Logs playbook_id on every outcome for win-rate analytics."),
    ]
    table = doc.add_table(rows=len(components) + 1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Module"
    table.rows[0].cells[1].text = "Responsibility"
    for i, (mod, resp) in enumerate(components, 1):
        table.rows[i].cells[0].text = mod
        table.rows[i].cells[1].text = resp

    doc.add_paragraph()
    doc.add_heading("4.3 Decision Rule Change", level=2)
    doc.add_paragraph("Today: IF confluence >= threshold AND gates pass → BUY.")
    doc.add_paragraph("Target: IF primary_playbook.trigger_fired AND gates pass → BUY with playbook_id.")


def add_state_machine(doc: Document) -> None:
    doc.add_heading("5. Play State Machine", level=1)

    states = [
        ("IDLE", "No eligible playbook or preconditions not met."),
        ("ARMED (WATCH)", "Preconditions satisfied; waiting for trigger event. Shown in Watch box."),
        ("TRIGGERED", "Trigger fired; pending final gate check before OPEN."),
        ("OPEN", "Position live. Shown in Open box."),
        ("MANAGING", "Trailing stops, scale rules, time stops active."),
        ("CLOSED", "Outcome recorded with playbook_id for telemetry."),
    ]
    table = doc.add_table(rows=len(states) + 1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "State"
    table.rows[0].cells[1].text = "Meaning"
    for i, (st, meaning) in enumerate(states, 1):
        table.rows[i].cells[0].text = st
        table.rows[i].cells[1].text = meaning

    doc.add_paragraph()
    doc.add_paragraph(
        "Only one PRIMARY armed playbook should be visible at a time. Secondary setups may "
        "appear as ranked alternates in a future iteration."
    )


def add_playbooks(doc: Document) -> None:
    doc.add_heading("6. The Twelve Playbooks (Full Catalog)", level=1)
    doc.add_paragraph(
        "Each playbook follows: Preconditions → Trigger → Invalidation → Target/Stop → "
        "Typical session window. IDs are stable for telemetry."
    )

    playbooks = [
        {
            "id": "PB-01",
            "name": "VWAP Reclaim",
            "direction": "Long or Short",
            "regime": "Trend / recovery after flush",
            "pre": "Price below VWAP ≥15m, then reclaims with volume; EMA9 curling toward VWAP.",
            "trigger": "Close above VWAP + hold 2 consecutive 3m bars; flow skew aligns.",
            "invalidate": "Close back below VWAP on volume; regime flips to chop.",
            "target": "Nearest call wall or prior high; stop below reclaim candle low.",
            "window": "09:45–14:00 ET",
        },
        {
            "id": "PB-02",
            "name": "VWAP Reject",
            "direction": "Short (or long inverse)",
            "regime": "Weak trend / distribution",
            "pre": "Rally into VWAP from below; repeated rejections at VWAP band.",
            "trigger": "3m close rejection wick + negative net flow spike.",
            "invalidate": "Acceptance above VWAP (2 closes).",
            "target": "Put wall / session low.",
            "window": "10:00–15:00 ET",
        },
        {
            "id": "PB-03",
            "name": "Opening Range Breakout (ORB)",
            "direction": "Break direction",
            "regime": "Opening drive",
            "pre": "First 15–30m range defined; GEX not pinning inside range.",
            "trigger": "Break of OR high/low with flow confirmation; spot clears flip level.",
            "invalidate": "Re-entry inside OR; halt feed degraded (optional strict mode).",
            "target": "1× OR width extension; wall beyond.",
            "window": "09:35–10:30 ET",
        },
        {
            "id": "PB-04",
            "name": "Gamma Pin Fade",
            "direction": "Fade toward pin",
            "regime": "High pin / low vol midday",
            "pre": "Spot between major walls; charm decay elevated; low ATR.",
            "trigger": "Touch of wall + rejection; confluence on mean-reversion factors.",
            "invalidate": "Sustained breakout through wall with flow.",
            "target": "Opposite wall or max pain.",
            "window": "11:30–15:00 ET",
        },
        {
            "id": "PB-05",
            "name": "Wall Break Continuation",
            "direction": "Break direction",
            "regime": "Trend / vol expansion",
            "pre": "Price compressed under call or put wall; rising VEX magnitude.",
            "trigger": "Close through wall + rising premium flow same direction.",
            "invalidate": "Immediate reclaim inside wall within 5m.",
            "target": "Next wall from matrix ladder.",
            "window": "10:00–15:30 ET",
        },
        {
            "id": "PB-06",
            "name": "Flip Level Ride",
            "direction": "With flip break",
            "regime": "Trend",
            "pre": "Spot oscillating at gamma flip; regime trending.",
            "trigger": "Decisive break of flip with EMA9/21 stack aligned.",
            "invalidate": "Recross flip and hold 3m.",
            "target": "Next flip or wall.",
            "window": "All RTH",
        },
        {
            "id": "PB-07",
            "name": "Max Pain Gravitation",
            "direction": "Toward max pain",
            "regime": "Expiry / pin",
            "pre": "Spot >0.3% from max pain; time >14:00; charm elevated.",
            "trigger": "Momentum stall toward pain; decreasing realized vol.",
            "invalidate": "Strong flow trend away from pain.",
            "target": "Max pain strike.",
            "window": "14:00–15:45 ET",
        },
        {
            "id": "PB-08",
            "name": "Power Hour Momentum",
            "direction": "Flow direction",
            "regime": "Power hour",
            "pre": "15:00–16:00; net flow dominant one side 10m+.",
            "trigger": "Break of 30m micro-range with accelerating prints.",
            "invalidate": "Flow flip + VWAP cross against.",
            "target": "Session H/L or wall.",
            "window": "15:00–15:55 ET",
        },
        {
            "id": "PB-09",
            "name": "HELIX Flow Surge",
            "direction": "Flow direction",
            "regime": "Any with premium spike",
            "pre": "HELIX alert tier ≥ threshold; ticker SPX/SPXW.",
            "trigger": "Desk direction aligns within 2 play polls; spot near strike cluster.",
            "invalidate": "No follow-through next poll; opposite surge.",
            "target": "Wall in surge direction.",
            "window": "All RTH",
        },
        {
            "id": "PB-10",
            "name": "EMA Stack Pullback",
            "direction": "Trend direction",
            "regime": "Trend",
            "pre": "EMA9 > EMA21 > SMA50 (bull) or inverse; pullback to EMA9/21.",
            "trigger": "Bounce candle + positive flow on 3m.",
            "invalidate": "Close through EMA21 against trend.",
            "target": "Prior swing / wall.",
            "window": "10:00–15:00 ET",
        },
        {
            "id": "PB-11",
            "name": "Range Chop Scalp",
            "direction": "Range fade",
            "regime": "Chop / low trend score",
            "pre": "Regime chop; defined 30m range; no breakout.",
            "trigger": "Fade at range edge with rejection wick.",
            "invalidate": "Range break with volume.",
            "target": "Mid-range or opposite edge.",
            "window": "11:00–14:00 ET",
        },
        {
            "id": "PB-12",
            "name": "Lotto Reversal",
            "direction": "Reversal",
            "regime": "Extreme extension",
            "pre": "Rapid extension >0.5% in 15m; RSI stretch; near wall.",
            "trigger": "Reversal candle + flow exhaustion signal.",
            "invalidate": "Continuation with new flow high.",
            "target": "VWAP or mid-range (tight stop).",
            "window": "All RTH (reduced size)",
        },
    ]

    for pb in playbooks:
        doc.add_heading(f"{pb['id']}: {pb['name']}", level=2)
        fields = [
            ("Direction", pb["direction"]),
            ("Regime tags", pb["regime"]),
            ("Preconditions", pb["pre"]),
            ("Trigger", pb["trigger"]),
            ("Invalidation", pb["invalidate"]),
            ("Target / Stop", pb["target"]),
            ("Session window", pb["window"]),
        ]
        for label, val in fields:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(val)


def add_data_plane(doc: Document) -> None:
    doc.add_heading("7. Cross-Tool Data Plane", level=1)
    doc.add_paragraph(
        "Playbooks must consume a unified data contract—not siloed scores. Sources:"
    )

    sources = [
        ("SpxDeskPayload", "Spot, session, regime, walls, flip, max pain, anchor, halt status."),
        ("PlayTechnicals", "3m/5m EMA/SMA/VWAP, breakouts, RSI, session H/L."),
        ("GEX matrix", "Strikes ladder, GEX/VEX/DEX/CHARM lenses, cross_validation vs UW."),
        ("Flows / HELIX", "Premium prints, net flow, alert tiers, persist path."),
        ("Night Hawk", "Bias nudge (±3)—advisory unless elevated to gate."),
        ("Largo", "Narration via getSpxPlayState—advisory only by default."),
    ]
    table = doc.add_table(rows=len(sources) + 1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Source"
    table.rows[0].cells[1].text = "Fields used by matcher"
    for i, (src, fields) in enumerate(sources, 1):
        table.rows[i].cells[0].text = src
        table.rows[i].cells[1].text = fields

    doc.add_paragraph()
    doc.add_heading("7.1 Minimum Contract Per Playbook Eval", level=2)
    for item in [
        "desk.spot, desk.regime, desk.gamma_flip, desk.walls, desk.max_pain",
        "technicals.vwap, ema9, ema21, sma50, breakout flags",
        "flows.net_flow_10m, helix.last_alert (if any)",
        "matrix.nearest_walls, cross_validation.status",
        "session.clock, is_rth, minutes_to_close",
    ]:
        doc.add_paragraph(item, style="List Bullet")


def add_confluence_vs_checklist(doc: Document) -> None:
    doc.add_heading("8. Confluence vs Playbook Checklists", level=1)
    doc.add_paragraph(
        "Today: ConfluenceFactorsPanel shows ~20 global factors with weights—informative but "
        "not tied to one setup."
    )
    doc.add_paragraph(
        "Target: When a playbook is ARMED, the Confluence panel becomes that playbook's "
        "checklist only. Example for PB-01 VWAP Reclaim:"
    )
    checks = [
        "☐ Below VWAP ≥15m (historical flag)",
        "☐ Reclaim candle closed above VWAP",
        "☐ 2 consecutive 3m holds",
        "☐ Flow skew bullish",
        "☐ Regime not chop",
    ]
    for c in checks:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_paragraph(
        "Global factor score may remain internal for matcher ranking but should not dominate "
        "the trader-facing UI."
    )


def add_parallel_engines(doc: Document) -> None:
    doc.add_heading("9. Largo, Night Hawk, and Parallel Engines", level=1)
    doc.add_paragraph(
        "Fragmentation today: main play engine, lotto path, power hour rules, Night Hawk ±3 "
        "nudge—all can influence perception without one primary playbook."
    )
    doc.add_heading("9.1 Consolidation Principle", level=2)
    for item in [
        "Lotto (PB-12) and Power Hour (PB-08) become registry entries, not separate silos.",
        "Night Hawk bias feeds regime router as input, not parallel BUY.",
        "Largo reads playbook_id + state for narration ('ARMED: VWAP Reclaim').",
    ]:
        doc.add_paragraph(item, style="List Bullet")


def add_ui_mapping(doc: Document) -> None:
    doc.add_heading("10. UI Mapping — Trade Alerts Panel", level=1)
    doc.add_paragraph(
        "Recent UI work (PR #722 merged): solid black panel, three regions—Open, Watch, Confluence."
    )

    mapping = [
        ("Open Play box", "State OPEN or MANAGING. Shows playbook name, direction, entry, stop, target."),
        ("Watch Play box", "State ARMED. Shows primary playbook name, trigger conditions remaining, countdown/window."),
        ("Confluence factors", "Checklist for armed playbook only (post-migration). Updates ~3s with play poll."),
    ]
    table = doc.add_table(rows=len(mapping) + 1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "UI Region"
    table.rows[0].cells[1].text = "Data binding"
    for i, (ui, bind) in enumerate(mapping, 1):
        table.rows[i].cells[0].text = ui
        table.rows[i].cells[1].text = bind

    doc.add_paragraph()
    doc.add_paragraph(
        "Header UI (PR #721): EMA/SMA/Session in top stats row with Regime/Flip/Max Pain; "
        "live spot with ▲/▼ above Play Engine. Halt degraded banner removed per user request."
    )


def add_telemetry(doc: Document) -> None:
    doc.add_heading("11. Telemetry and Win-Rate Measurement", level=1)
    doc.add_paragraph(
        "Today: cold_buy vs watch_promote win rates exist; no playbook_id on outcomes."
    )
    doc.add_heading("11.1 Required Fields (Proposed)", level=2)
    fields = [
        "playbook_id (e.g. PB-01)",
        "playbook_state transitions with timestamps",
        "trigger_reason (which condition fired)",
        "invalidation_reason (if closed early)",
        "outcome: win | loss | scratch",
        "hold_duration_sec",
        "regime_at_entry",
    ]
    for f in fields:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_paragraph(
        "Enables: 'PB-03 ORB win rate 62% last 20 sessions' instead of aggregate confluence buckets."
    )


def add_migration(doc: Document) -> None:
    doc.add_heading("12. Migration Roadmap (Shadow → Live)", level=1)

    phases = [
        ("Phase 0 — Document", "This spec + PLAYBOOK-REGISTRY.md in repo (canonical)."),
        ("Phase 1 — Shadow", "Matcher runs alongside legacy engine; log would-be playbook_id without changing BUY."),
        ("Phase 2 — ARM UI", "Watch box shows primary playbook from shadow matcher."),
        ("Phase 3 — Trigger gate", "BUY requires playbook trigger; legacy score demoted to tie-breaker."),
        ("Phase 4 — Telemetry", "playbook_id on all outcomes; dashboard per-pattern win rates."),
        ("Phase 5 — Deprecate", "Remove coarse 0dte watch key; remove duplicate lotto/power-hour silos."),
    ]
    for i, (phase, desc) in enumerate(phases, 1):
        p = doc.add_paragraph()
        p.add_run(f"{phase}: ").bold = True
        p.add_run(desc)


def add_open_decisions(doc: Document) -> None:
    doc.add_heading("13. Open Design Decisions", level=1)
    decisions = [
        "Which 3 playbooks are true A+ trades (highest size)?",
        "Lunch 11:30–13:30: flat (no new arms) vs pin-only (PB-04, PB-07)?",
        "One open play maximum—still yes?",
        "Starter size vs full-only entries?",
        "Claude/Largo: advisory only vs hard gate on low confidence?",
        "Play poll 3s vs 1s for trigger detection?",
        "Halt feed degraded: fail-open (today) vs fail-closed for ORB/breakout playbooks?",
    ]
    for d in decisions:
        doc.add_paragraph(d, style="List Bullet")


def add_validation_checklist(doc: Document) -> None:
    doc.add_heading("14. Validation Checklist for Other AIs", level=1)
    doc.add_paragraph(
        "When sharing this document with Claude, ChatGPT, or reviewers, ask them to evaluate:"
    )

    questions = [
        "Is one primary armed playbook sufficient, or should 2–3 ranked watches be shown?",
        "Are twelve playbooks too many for 0DTE SPX? Which should merge or drop?",
        "Are preconditions/triggers measurable with existing desk fields, or what gaps exist?",
        "Does the state machine miss states (e.g. COOLDOWN after invalidation)?",
        "Is shadow-mode migration the right risk profile?",
        "What failure modes occur if regime router misclassifies chop as trend?",
        "How should playbook conflicts resolve (e.g. PB-05 wall break vs PB-04 pin fade)?",
        "Suggested priority order for Phase 1–3 implementation.",
    ]
    for q in questions:
        doc.add_paragraph(q, style="List Number")

    doc.add_paragraph()
    doc.add_paragraph(
        "Prompt template for other AIs: 'You are reviewing a trading-system design doc for "
        "0DTE SPX options. Critique the playbook-first architecture in Section 4–6. Identify "
        "gaps, overlaps, and suggest the top 3 playbooks to implement first with measurable "
        "triggers using the data in Section 7.'"
    )


def add_glossary(doc: Document) -> None:
    doc.add_heading("15. Glossary", level=1)
    terms = [
        ("ARMED", "Playbook preconditions met; waiting for trigger (WATCH state)."),
        ("Confluence", "Weighted sum of factor scores (legacy); becoming playbook checklist."),
        ("Desk", "SpxDeskPayload—canonical per-tick context for SPX Slayer."),
        ("Flip", "Gamma flip level—dealers delta-neutral pivot."),
        ("GEX", "Gamma exposure by strike; walls where dealer hedging intensifies."),
        ("HELIX", "Flow alert surface; high-premium option prints."),
        ("Largo", "AI narration layer; reads play state, does not gate."),
        ("Matcher", "Selects best eligible playbook from registry."),
        ("0DTE", "Zero days to expiration—same-day SPX/SPXW options."),
        ("playbook_id", "Stable identifier (PB-01..PB-12) for telemetry."),
        ("Regime Router", "Classifies market context to filter eligible playbooks."),
        ("RTH", "Regular trading hours 09:30–16:00 ET."),
        ("Trigger", "Event that converts ARMED → OPEN (entry signal)."),
        ("Watch key", "Legacy coarse id; replaced by playbook instance id."),
    ]
    table = doc.add_table(rows=len(terms) + 1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Definition"
    for i, (term, defn) in enumerate(terms, 1):
        table.rows[i].cells[0].text = term
        table.rows[i].cells[1].text = defn


def add_footer(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Document Control", level=1)
    doc.add_paragraph("Author: BlackOut / Cursor Cloud Agent")
    doc.add_paragraph(f"Generated: {datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("Status: Design phase — not yet implemented in production code.")
    doc.add_paragraph(
        "Related PRs: #722 Trade Alerts UI (merged); #721 Header/halt banner (pending merge at time of writing)."
    )
    doc.add_paragraph(
        "Canonical code paths: src/features/spx/, evaluateSpxPlay, computeSpxConfluence, useSpxPlay.ts"
    )


def main() -> None:
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc = Document()
    set_doc_styles(doc)
    add_title_page(doc)
    add_toc_placeholder(doc)
    add_executive_summary(doc)
    add_problem(doc)
    add_current_arch(doc)
    add_target_arch(doc)
    add_state_machine(doc)
    add_playbooks(doc)
    add_data_plane(doc)
    add_confluence_vs_checklist(doc)
    add_parallel_engines(doc)
    add_ui_mapping(doc)
    add_telemetry(doc)
    add_migration(doc)
    add_open_decisions(doc)
    add_validation_checklist(doc)
    add_glossary(doc)
    add_footer(doc)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")
    print(f"Size: {os.path.getsize(OUT_PATH)} bytes")


if __name__ == "__main__":
    main()
